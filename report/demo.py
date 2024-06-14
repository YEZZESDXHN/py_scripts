import re

import pandas as pd
from docx import Document


# def replace_placeholder(doc, params):
#     """替换占位符"""
#     for paragraph in doc.paragraphs:
#         for param in params:
#             pv = str(params[param])
#             # ph = f'ph_{param}'
#             if param in paragraph.text:
#                 for run in paragraph.runs:
#                     if param in run.text:
#                         run.text = run.text.replace(param, pv)
#                         run.italic = False

def replace_word(doc, params):
    """
    定义批量替换文字的函数
    :param doc: 要替换的文档
    :param old_word: 被替换的文字
    :param new_word: 替换后的文字
    :return:
    """
    for param in params:
        for paragraph in doc.paragraphs:
            if param in paragraph.text:
                tmp = ''
                runs = paragraph.runs
                for i, run in enumerate(runs):
                    tmp += run.text  # 合并 run 字符串
                if param in tmp:
                    # 如果存在匹配的字符串，将当前 run 替换为合并后的字符串
                    run.text = run.text.replace(run.text, tmp)
                    run.text = run.text.replace(param, params[param])
                    tmp = ''
                else:
                    # 如果没有匹配到目标字符串，将当前 run 置空
                    run.text = run.text.replace(run.text, '')
                    if i == len(runs) - 1:
                        # 如果是当前段落一直没有符合规则的字符串，直接将当前 run 替换为 tmp
                        run.add_text(tmp)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if param in cell.text:
                        cell.text = cell.text.replace(param, params[param])





class rtm():
    def __init__(self):
        self.total_case = None
        self.total_case_pass = None
        self.total_case_fail = None
        self.total_case_NT = None
        self.total_case_NA = None
        self.total_case_test_result_Undefined = None
        self.total_case_Automatic = None
        self.total_case_Semi_Automatic = None
        self.total_case_Manual = None
        self.total_case_test_type_Undefined = None

        self.total_req = None
        self.total_req_pass = None
        self.total_req_fail = None
        self.total_req_Partially_Failed = None
        self.total_req_Partially_Passed = None
        self.total_req_NT = None
        self.total_req_NA = None
        self.total_req_test_result_Undefined = None

        self.total_module_num = None
        self.req_module_summary = None
        self.case_module_summary = None

        self.rtm_raw_df = None
        self.rtm_init_df = None
        self.module_num = None
        self.req_id_col_name = 'ID'
        self.req_module_col_name = 'Module'
        self.req_adASIL_col_name = 'aASIL'
        self.req_result_col_name = 'Consolidated Test Result'
        self.case_id_col_name = 'ID.1'
        self.case_module_col_name = 'Module.1'
        self.case_adASIL_col_name = 'adASIL'
        self.case_result_col_name = 'Test Result'
        self.case_TestType_col_name = 'aTestType'

    def read_excel(self, *args, **kwargs):

        try:
            # 使用 *args 和 **kwargs 调用 pd.read_excel
            self.rtm_raw_df = pd.read_excel(*args, **kwargs)
            self.rtm_init_df = self.rtm_raw_df.loc[:,
                               [self.req_module_col_name, self.req_id_col_name, self.req_adASIL_col_name,
                                self.req_result_col_name,
                                self.case_module_col_name, self.case_id_col_name, self.case_adASIL_col_name,
                                self.case_result_col_name, self.case_TestType_col_name]]
            # self.rtm_init_df['req_mouule_id']=self.rtm_init_df[self.req_module_col_name]+self.rtm_init_df[self.req_id_col_name]

            # 在id列后面增加一列req_module_id
            self.rtm_init_df.insert(self.rtm_init_df.columns.get_loc(self.req_id_col_name) + 1, 'req_module_id',
                                    self.rtm_init_df[self.req_module_col_name] + self.rtm_init_df[self.req_id_col_name])
            self.rtm_init_df.insert(self.rtm_init_df.columns.get_loc(self.case_id_col_name) + 1, 'case_module_id',
                                    self.rtm_init_df[self.case_module_col_name] + self.rtm_init_df[
                                        self.case_id_col_name])
            self.rtm_init_df.to_excel('./output/cal_out.xlsx', sheet_name='init')
        except Exception as e:
            print(f"Error reading Excel file: {e}")



    def generate_report(self):
        # 测试情况填入world,生成报告
        res_doc = './output/report.docx'
        params = {
            "{{case_Pl}}": 'self.total_case',
        }
        tpl = Document('./tpl/jx/CDA_DFC_Project System Integration Test Report - tpl.docx')
        replace_word(tpl, params)
        # replace_word(tpl,'{{case_Pl}}',str(self.total_case))
        tpl.save(res_doc)


    # 汇总分析测试情况
    def summary_module(self):
        # 提取req和case的module列去除空行生成Series req_module，case_module
        req_module = self.rtm_init_df[self.req_module_col_name].dropna().drop_duplicates().reset_index(drop=True)
        case_module = self.rtm_init_df[self.case_module_col_name].dropna().drop_duplicates().reset_index(drop=True)
        num1 = req_module.size
        num2 = case_module.size
        if num1 == num2:
            self.total_module_num = num1

            # 将module列生成的Series转为dataframe，self.req_module_summary后续存放汇总信息
            self.req_module_summary = req_module.to_frame()
            # 向汇总df中添加一列使用正则删除原始module中数据前面的路径信息
            self.req_module_summary['module_short'] = self.req_module_summary[self.req_module_col_name].apply(
                lambda x: re.search(r'\d{4}_(\w+)$', x).group(0) if re.search(r'\d{4}_(\w+)$', x) else None
            )
            # 添加空列Pass_num
            self.req_module_summary['total'] = None
            self.req_module_summary['Pass'] = None
            self.req_module_summary['Fail'] = None
            self.req_module_summary['Partially_Failed'] = None
            self.req_module_summary['Partially_Passed'] = None
            self.req_module_summary['NT'] = None
            self.req_module_summary['NA'] = None
            self.req_module_summary['test_result_Undefined'] = None
            self.req_module_summary['pass_rate'] = None

            self.case_module_summary = case_module.to_frame()
            self.case_module_summary['module_short'] = self.case_module_summary[self.case_module_col_name].apply(
                lambda x: re.search(r'\d{4}_(\w+)$', x).group(0) if re.search(r'\d{4}_(\w+)$', x) else None
            )
            self.case_module_summary['total'] = None
            self.case_module_summary['Pass'] = None
            self.case_module_summary['Fail'] = None
            self.case_module_summary['NT'] = None
            self.case_module_summary['NA'] = None
            self.case_module_summary['test_result_Undefined'] = None
            self.case_module_summary['pass_rate'] = None

            self.case_module_summary['Automatic'] = None
            self.case_module_summary['Semi-Automatic'] = None
            self.case_module_summary['Manual'] = None
            self.case_module_summary['test_type_Undefined'] = None
            self.case_module_summary['Automatic_rate'] = None

            req_module_list = []
            for i in range(num1):
                # 按模块生成dataframe，去除空行存放到list中
                req_module_list.append(self.rtm_init_df.loc[
                                           self.rtm_init_df[self.req_module_col_name] == req_module[i], [
                                               self.req_module_col_name, self.req_id_col_name, 'req_module_id',
                                               self.req_adASIL_col_name, self.req_result_col_name]])
                # drop_duplicates去重函数
                # subset - - 指定特定的列 默认所有列
                # keep: {'first', 'last', False} - - 删除重复项并保留第一次出现的项 默认第一个
                # keep = False - - 表示删除所有重复项 不保留
                # inplace - - 是否直接修改原对象
                # ignore_index = True - - 重置索引 （version 1.0.0 才有这个参数）
                req_module_list[i].drop_duplicates(subset="req_module_id", keep='first', inplace=True)
                # self.req_module_summary.loc[self.req_module_summary["Module"]==req_module_list[i]["Module"][1],'Pass']=req_module_list[i].

                # 写入汇总页各模块执行情况
                try:
                    self.req_module_summary.loc[i, 'total'] = req_module_list[i].shape[0]
                except:
                    self.req_module_summary.loc[i, 'total'] = 0

                try:
                    self.req_module_summary.loc[i, 'Pass'] = \
                    req_module_list[i][self.req_result_col_name].value_counts()['Passed']

                except:
                    self.req_module_summary.loc[i, 'Pass'] = 0

                try:
                    self.req_module_summary.loc[i, 'Fail'] = \
                        req_module_list[i][self.req_result_col_name].value_counts()['Failed']

                except:
                    self.req_module_summary.loc[i, 'Fail'] = 0

                try:
                    self.req_module_summary.loc[i, 'Partially_Failed'] = \
                        req_module_list[i][self.req_result_col_name].value_counts()['Partially Failed']

                except:
                    self.req_module_summary.loc[i, 'Partially_Failed'] = 0

                try:
                    self.req_module_summary.loc[i, 'Partially_Passed'] = \
                        req_module_list[i][self.req_result_col_name].value_counts()['Partially Passed']

                except:
                    self.req_module_summary.loc[i, 'Partially_Passed'] = 0

                try:
                    self.req_module_summary.loc[i, 'NT'] = \
                        req_module_list[i][self.req_result_col_name].value_counts()['Not Executed']

                except:
                    self.req_module_summary.loc[i, 'NT'] = 0

                try:
                    self.req_module_summary.loc[i, 'NA'] = \
                        req_module_list[i][self.req_result_col_name].value_counts()['Not Executable']

                except:
                    self.req_module_summary.loc[i, 'NA'] = 0

                try:
                    self.req_module_summary.loc[i, 'test_result_Undefined'] = \
                        req_module_list[i][self.req_result_col_name].value_counts()['n/a']

                except:
                    self.req_module_summary.loc[i, 'test_result_Undefined'] = 0

                try:
                    self.req_module_summary.loc[i, 'pass_rate'] = self.req_module_summary.loc[i, 'Pass'] / \
                                                                  self.req_module_summary.loc[i, 'total']
                except:
                    self.req_module_summary.loc[i, 'pass_rate'] = "error"

            self.req_module_summary.loc[num1] = ['/', '/',
                                                 self.total_req,
                                                 self.total_req_pass,
                                                 self.total_req_fail,
                                                 self.total_req_Partially_Failed,
                                                 self.total_req_Partially_Passed,
                                                 self.total_req_NT,
                                                 self.total_req_NA,
                                                 self.total_req_test_result_Undefined,
                                                 'Nan']
            try:
                self.req_module_summary.loc[num1, 'pass_rate'] = self.req_module_summary.loc[num1, 'Pass'] / \
                                                                 self.req_module_summary.loc[num1, 'total']
            except:
                self.req_module_summary.loc[num1, 'pass_rate'] = "error"

            case_module_list = []
            for i in range(num1):
                # 按模块生成dataframe，去除空行存放到list中
                case_module_list.append(self.rtm_init_df.loc[
                                            self.rtm_init_df[self.case_module_col_name] == case_module[i], [
                                                self.case_module_col_name, self.case_id_col_name, 'case_module_id',
                                                self.case_adASIL_col_name, self.case_result_col_name,
                                                self.case_TestType_col_name]])
                # drop_duplicates去重函数
                case_module_list[i].drop_duplicates(subset="case_module_id", keep='first', inplace=True)

                # 写入汇总页各模块执行情况
                try:
                    self.case_module_summary.loc[i, 'total'] = case_module_list[i].shape[0]
                except:
                    self.case_module_summary.loc[i, 'total'] = 0

                try:
                    self.case_module_summary.loc[i, 'Pass'] = \
                        case_module_list[i][self.case_result_col_name].value_counts()['Passed']

                except:
                    self.case_module_summary.loc[i, 'Pass'] = 0

                try:
                    self.case_module_summary.loc[i, 'Fail'] = \
                        case_module_list[i][self.case_result_col_name].value_counts()['Failed']

                except:
                    self.case_module_summary.loc[i, 'Fail'] = 0

                try:
                    self.case_module_summary.loc[i, 'NT'] = \
                        case_module_list[i][self.case_result_col_name].value_counts()['Not Executed']

                except:
                    self.case_module_summary.loc[i, 'NT'] = 0

                try:
                    self.case_module_summary.loc[i, 'NA'] = \
                        case_module_list[i][self.case_result_col_name].value_counts()['Not Executable']

                except:
                    self.case_module_summary.loc[i, 'NA'] = 0

                try:
                    self.case_module_summary.loc[i, 'test_result_Undefined'] = \
                        case_module_list[i][self.case_result_col_name].value_counts()['n/a']

                except:
                    self.case_module_summary.loc[i, 'test_result_Undefined'] = 0

                try:
                    self.case_module_summary.loc[i, 'pass_rate'] = self.case_module_summary.loc[i, 'Pass'] / \
                                                                   self.case_module_summary.loc[i, 'total']
                except:
                    self.case_module_summary.loc[i, 'pass_rate'] = "error"

                try:
                    self.case_module_summary.loc[i, 'Automatic'] = \
                        case_module_list[i][self.case_TestType_col_name].value_counts()['Automatic']

                except:
                    self.case_module_summary.loc[i, 'Automatic'] = 0

                try:
                    self.case_module_summary.loc[i, 'Semi-Automatic'] = \
                        case_module_list[i][self.case_TestType_col_name].value_counts()['Semi-Automatic']

                except:
                    self.case_module_summary.loc[i, 'Semi-Automatic'] = 0

                try:
                    self.case_module_summary.loc[i, 'Manual'] = \
                        case_module_list[i][self.case_TestType_col_name].value_counts()['Manual']

                except:
                    self.case_module_summary.loc[i, 'Manual'] = 0

                try:
                    self.case_module_summary.loc[i, 'test_type_Undefined'] = \
                        case_module_list[i][self.case_TestType_col_name].value_counts()['n/a']

                except:
                    self.case_module_summary.loc[i, 'test_type_Undefined'] = 0

                try:
                    self.case_module_summary.loc[i, 'Automatic_rate'] = (self.case_module_summary.loc[i, 'Automatic'] +
                                                                         self.case_module_summary.loc[
                                                                             i, 'Semi-Automatic']) / \
                                                                        self.case_module_summary.loc[i, 'total']
                except:
                    self.case_module_summary.loc[i, 'Automatic_rate'] = "error"

            self.case_module_summary.loc[num1] = ['/', '/',
                                                  self.total_case,
                                                  self.total_case_pass,
                                                  self.total_case_fail,
                                                  self.total_case_NT,
                                                  self.total_case_NA,
                                                  self.total_case_test_result_Undefined,
                                                  'Nan',
                                                  self.total_case_Automatic,
                                                  self.total_case_Semi_Automatic,
                                                  self.total_case_Manual,
                                                  self.total_case_test_type_Undefined,
                                                  'Nan']
            try:
                self.case_module_summary.loc[num1, 'Automatic_rate'] = (self.case_module_summary.loc[
                                                                            num1, 'Automatic'] +
                                                                        self.case_module_summary.loc[
                                                                            num1, 'Semi-Automatic']) / \
                                                                       self.case_module_summary.loc[num1, 'total']
            except:
                self.case_module_summary.loc[num1, 'Automatic_rate'] = "error"

            try:
                self.case_module_summary.loc[num1, 'pass_rate'] = self.case_module_summary.loc[num1, 'Pass'] / \
                                                                  self.case_module_summary.loc[num1, 'total']
            except:
                self.case_module_summary.loc[num1, 'pass_rate'] = "error"

            try:
                # 向excel中追加工作表
                with pd.ExcelWriter("./output/cal_out.xlsx", mode='a', engine='openpyxl') as writer:
                    self.req_module_summary.to_excel(writer, sheet_name='req_module_summary')
                    self.case_module_summary.to_excel(writer, sheet_name='case_module_summary')

                    # # 将req_module_list和case_module_list写入excel，仅作调试使用
                    # for i in range(num1):
                    #     req_module_list[i].to_excel(writer,sheet_name="req"+str(i))
                    # for i in range(num1):
                    #     case_module_list[i].to_excel(writer, sheet_name="case"+str(i))
            except Exception as e:
                print(f"Error to Excel sheet: {e}")

    def data_analysis(self):
        # 需case总体测试情况
        self.total_case = self.rtm_init_df['case_module_id'].nunique()
        # drop_duplicates去重函数
        # subset - - 指定特定的列 默认所有列
        # keep: {'first', 'last', False} - - 删除重复项并保留第一次出现的项 默认第一个
        # keep = False - - 表示删除所有重复项 不保留
        # inplace - - 是否直接修改原对象
        # ignore_index = True - - 重置索引 （version 1.0.0 才有这个参数）
        case_clear_rtm_init_df = self.rtm_init_df.drop_duplicates(subset="case_module_id", keep='first', inplace=False)
        try:
            self.total_case_pass = case_clear_rtm_init_df[self.case_result_col_name].value_counts()['Passed']
        except:
            self.total_case_pass = 0

        try:
            self.total_case_fail = case_clear_rtm_init_df[self.case_result_col_name].value_counts()['Failed']
        except:
            self.total_case_fail = 0
        try:
            self.total_case_NT = case_clear_rtm_init_df[self.case_result_col_name].value_counts()['Not Executed']
        except:
            self.total_case_NT = 0

        try:
            self.total_case_NA = case_clear_rtm_init_df[self.case_result_col_name].value_counts()['Not Executable']
        except:
            self.total_case_NA = 0

        try:
            self.total_case_test_result_Undefined = case_clear_rtm_init_df[self.case_result_col_name].value_counts()[
                'n/a']
        except:
            self.total_case_test_result_Undefined = 0

        try:
            self.total_case_Automatic = case_clear_rtm_init_df[self.case_TestType_col_name].value_counts()['Automatic']
        except:
            self.total_case_Automatic = 0
        try:
            self.total_case_Semi_Automatic = case_clear_rtm_init_df[self.case_TestType_col_name].value_counts()[
                'Semi-Automatic']
        except:
            self.total_case_Semi_Automatic = 0

        try:
            self.total_case_Manual = case_clear_rtm_init_df[self.case_TestType_col_name].value_counts()['Manual']
        except:
            self.total_case_Manual = 0

        try:
            self.total_case_test_type_Undefined = case_clear_rtm_init_df[self.case_TestType_col_name].value_counts()[
                'n/a']
        except:
            self.total_case_test_type_Undefined = 0

        # 需求总体测试情况
        self.total_req = self.rtm_init_df['req_module_id'].nunique()

        req_clear_rtm_init_df = self.rtm_init_df.drop_duplicates(subset="req_module_id", keep='first', inplace=False)

        try:
            self.total_req_pass = req_clear_rtm_init_df[self.req_result_col_name].value_counts()['Passed']
        except:
            self.total_req_pass = 0

        try:
            self.total_req_fail = req_clear_rtm_init_df[self.req_result_col_name].value_counts()['Failed']
        except:
            self.total_req_fail = 0
        try:
            self.total_req_NT = req_clear_rtm_init_df[self.req_result_col_name].value_counts()['Not Executed']
        except:
            self.total_req_NT = 0

        try:
            self.total_req_NA = req_clear_rtm_init_df[self.req_result_col_name].value_counts()['Not Executable']
        except:
            self.total_req_NA = 0

        try:
            self.total_req_test_result_Undefined = req_clear_rtm_init_df[self.req_result_col_name].value_counts()['n/a']
        except:
            self.total_req_test_result_Undefined = 0

        try:
            self.total_req_Partially_Failed = req_clear_rtm_init_df[self.req_result_col_name].value_counts()[
                'Partially Failed']
        except:
            self.total_req_Partially_Failed = 0

        try:
            self.total_req_Partially_Passed = req_clear_rtm_init_df[self.req_result_col_name].value_counts()[
                'Partially Passed']
        except:
            self.total_req_Partially_Passed = 0


        #测试情况汇总
        self.summary_module()

        #生成报告
        self.generate_report()


if __name__ == "__main__":
    rtm1 = rtm()
    rtm1.read_excel(io='./RTM/RTM.xlsx', sheet_name='Sysintg', skiprows=1)
    rtm1.data_analysis()
